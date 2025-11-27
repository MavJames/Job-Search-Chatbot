"""
MCP Server for creating properly formatted Word documents (.docx)
"""

import sys
import traceback

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastmcp import FastMCP

# ============================================================================
# Theme Configuration
# ============================================================================

THEME = {
    "font_color": RGBColor(0, 0, 0),
    "resume": {
        "margin": 0.5,
        "name_size": Pt(16),
        "contact_size": Pt(9),
        "contact_spacing": Pt(6),
        "section_heading_size": Pt(11),
        "section_space_before": Pt(6),
        "section_space_after": Pt(4),
        "summary_size": Pt(10),
        "summary_spacing": Pt(6),
        "job_title_size": Pt(10),
        "job_details_size": Pt(9),
        "bullet_size": Pt(10),
        "education_degree_size": Pt(10),
        "education_school_size": Pt(9),
        "skills_size": Pt(10),
    },
    "cover_letter": {
        "margin": 1.0,
        "font_name": "Arial",
        "name_size": Pt(24),
        "title_size": Pt(14),
        "contact_size": Pt(10),
        "date_size": Pt(10),
        "recipient_size": Pt(10),
        "reference_size": Pt(10),
        "body_size": Pt(10),
        "closing_size": Pt(10),
        "signature_size": Pt(14),
        "line_spacing": 1.15,
        "space_after_contact": Pt(12),
        "space_after_date": Pt(12),
        "space_after_recipient": Pt(12),
        "space_after_reference": Pt(12),
        "space_after_body": Pt(12),
        "space_after_closing": Pt(24),
    },
}

# ============================================================================
# Initialization
# ============================================================================

# Initialize FastMCP server
mcp = FastMCP("docx-creator")

# ============================================================================
# Helper Functions
# ============================================================================


def add_section_heading(doc: Document, text: str):
    """Add a formatted section heading (compact for one-page layout)"""
    cfg = THEME["resume"]
    heading = doc.add_paragraph(text)
    heading_run = heading.runs[0]
    heading_run.font.size = cfg["section_heading_size"]
    heading_run.font.bold = True
    heading_run.font.color.rgb = THEME["font_color"]
    heading.paragraph_format.space_before = cfg["section_space_before"]
    heading.paragraph_format.space_after = cfg["section_space_after"]


def _setup_page_margins(doc: Document, margin_inches: float):
    """Set page margins"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)


def _set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def _add_contact_info(doc: Document, contact: dict):
    """Add contact info section"""
    if not contact:
        return

    cfg = THEME["resume"]
    contact_parts = []
    if "email" in contact:
        contact_parts.append(contact["email"])
    if "phone" in contact:
        contact_parts.append(contact["phone"])
    if "location" in contact:
        contact_parts.append(contact["location"])
    if "linkedin" in contact:
        contact_parts.append(contact["linkedin"])

    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = cfg["contact_size"]
        contact_para.paragraph_format.space_after = cfg["contact_spacing"]


def _add_experience_section(doc: Document, experience: list):
    """Add experience section"""
    if not experience:
        return

    cfg = THEME["resume"]
    add_section_heading(doc, "EXPERIENCE")
    for i, exp in enumerate(experience):
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_after = Pt(1)
        title_run = title_para.add_run(
            f"{exp.get('title', '')} - {exp.get('company', '')}"
        )
        title_run.bold = True
        title_run.font.size = cfg["job_title_size"]

        details_para = doc.add_paragraph()
        details_para.paragraph_format.space_after = Pt(2)
        details_run = details_para.add_run(
            f"{exp.get('location', '')} | {exp.get('dates', '')}"
        )
        details_run.italic = True
        details_run.font.size = cfg["job_details_size"]

        if "responsibilities" in exp:
            for resp in exp["responsibilities"]:
                bullet_para = doc.add_paragraph(resp, style="List Bullet")
                bullet_para.paragraph_format.space_after = Pt(1)
                bullet_para.paragraph_format.line_spacing = 1.0
                for run in bullet_para.runs:
                    run.font.size = cfg["bullet_size"]

        if i < len(experience) - 1:
            space_para = doc.add_paragraph()
            space_para.paragraph_format.space_after = Pt(4)


def _add_education_section(doc: Document, education: list):
    """Add education section"""
    if not education:
        return

    cfg = THEME["resume"]
    add_section_heading(doc, "EDUCATION")
    for edu in education:
        edu_para = doc.add_paragraph()
        edu_para.paragraph_format.space_after = Pt(1)
        degree_run = edu_para.add_run(edu.get("degree", ""))
        degree_run.bold = True
        degree_run.font.size = cfg["education_degree_size"]

        school_para = doc.add_paragraph(
            f"{edu.get('school', '')} - {edu.get('location', '')} | {edu.get('graduation', '')}"
        )
        school_para.paragraph_format.space_after = Pt(4)
        school_para.runs[0].font.size = cfg["education_school_size"]


# ============================================================================
# MCP Tools
# ============================================================================


@mcp.tool()
def create_resume(
    output_path: str,
    name: str,
    contact: dict = None,
    summary: str = None,
    experience: list = None,
    education: list = None,
    skills: list = None,
) -> str:
    """
    Create a professionally formatted resume in .docx format (optimized for one page)
    """
    try:
        cfg = THEME["resume"]
        doc = Document()
        _setup_page_margins(doc, cfg["margin"])

        name_para = doc.add_paragraph(name)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.runs[0]
        name_run.font.size = cfg["name_size"]
        name_run.font.bold = True
        name_para.paragraph_format.space_after = Pt(2)

        _add_contact_info(doc, contact)

        if summary:
            add_section_heading(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph(summary)
            summary_para.paragraph_format.space_after = cfg["summary_spacing"]
            for run in summary_para.runs:
                run.font.size = cfg["summary_size"]

        _add_education_section(doc, education)
        _add_experience_section(doc, experience)

        if skills:
            add_section_heading(doc, "SKILLS")
            skills_para = doc.add_paragraph(", ".join(skills))
            for run in skills_para.runs:
                run.font.size = cfg["skills_size"]

        doc.save(output_path)
        return f"Resume successfully created at: {output_path}"

    except Exception as e:
        raise Exception(f"Error creating resume: {str(e)}")


@mcp.tool()
def create_cover_letter(
    output_path: str,
    name: str,
    body_paragraphs: list,
    contact: dict = None,
    date: str = None,
    recipient: dict = None,
    job_title: str = None,
    job_reference: str = None,
) -> str:
    """
    Create a professionally formatted cover letter in .docx format.
    Matches the "Olivia Wilson" design style.
    """
    try:
        cfg = THEME["cover_letter"]
        doc = Document()
        _setup_page_margins(doc, cfg["margin"])

        # Set default font
        style = doc.styles['Normal']
        style.font.name = cfg["font_name"]
        style.font.size = cfg["body_size"]

        # ====================================================================
        # Header Section (Table with 2 columns)
        # ====================================================================
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.allow_autofit = False
        
        # Set column widths (approximate for 8.5x11 with 1" margins = 6.5" printable)
        # We'll split 50/50
        for cell in header_table.rows[0].cells:
            cell.width = Inches(3.25)

        # Left Cell: Name and Title
        left_cell = header_table.cell(0, 0)
        # Clear default paragraph
        left_cell._element.clear_content()
        
        # Name
        name_para = left_cell.add_paragraph()
        name_run = name_para.add_run(name.upper())
        name_run.font.size = cfg["name_size"]
        name_run.font.bold = False  # The image shows it regular but large, maybe slightly bold? Let's stick to regular but large.
        # Actually image looks slightly bold or just tracking. Let's do normal.
        name_para.paragraph_format.space_after = Pt(0)
        
        # Job Title (if provided)
        if job_title:
            title_para = left_cell.add_paragraph()
            title_run = title_para.add_run(job_title)
            title_run.font.size = cfg["title_size"]
            title_run.font.bold = True # Image looks bold
            title_para.paragraph_format.space_after = Pt(0)

        # Right Cell: Contact Info
        right_cell = header_table.cell(0, 1)
        right_cell._element.clear_content()
        
        if contact:
            # Phone
            if "phone" in contact:
                p = right_cell.add_paragraph(contact["phone"])
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.runs[0].font.size = cfg["contact_size"]
                p.paragraph_format.space_after = Pt(0)
            # Email
            if "email" in contact:
                p = right_cell.add_paragraph(contact["email"])
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.runs[0].font.size = cfg["contact_size"]
                p.paragraph_format.space_after = Pt(0)
            # Address/Location
            if "address" in contact:
                p = right_cell.add_paragraph(contact["address"])
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.runs[0].font.size = cfg["contact_size"]
                p.paragraph_format.space_after = Pt(0)
            elif "location" in contact:
                p = right_cell.add_paragraph(contact["location"])
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.runs[0].font.size = cfg["contact_size"]
                p.paragraph_format.space_after = Pt(0)

        # Add bottom border to the header cells to create the line
        for cell in header_table.rows[0].cells:
            _set_cell_border(
                cell, 
                bottom={"sz": 6, "color": "000000", "val": "single"}
            )

        # Add some space after the header line
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

        # ====================================================================
        # Date (Right Aligned)
        # ====================================================================
        if date:
            date_para = doc.add_paragraph(date)
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.runs[0].font.size = cfg["date_size"]
            date_para.paragraph_format.space_after = cfg["space_after_date"]

        # ====================================================================
        # Recipient Info (Left Aligned)
        # ====================================================================
        if recipient:
            if "name" in recipient:
                p = doc.add_paragraph(recipient["name"])
                p.paragraph_format.space_after = Pt(0)
            if "title" in recipient:
                p = doc.add_paragraph(recipient["title"])
                p.paragraph_format.space_after = Pt(0)
            if "company" in recipient:
                p = doc.add_paragraph(recipient["company"])
                p.paragraph_format.space_after = Pt(0)
            if "address" in recipient:
                p = doc.add_paragraph(recipient["address"])
                p.paragraph_format.space_after = Pt(0)
            
            # Add space after recipient block
            doc.add_paragraph().paragraph_format.space_after = cfg["space_after_recipient"]

        # ====================================================================
        # Job Reference (Bold, Uppercase)
        # ====================================================================
        if job_reference:
            ref_para = doc.add_paragraph()
            ref_run = ref_para.add_run(f"JOB REFERENCE: {job_reference.upper()}")
            ref_run.font.bold = True
            ref_run.font.size = cfg["reference_size"]
            ref_run.font.name = cfg["font_name"]
            # Add extra spacing
            ref_para.paragraph_format.space_before = Pt(12)
            ref_para.paragraph_format.space_after = cfg["space_after_reference"]

        # ====================================================================
        # Salutation
        # ====================================================================
        # Try to find a name to address
        salutation_name = "Hiring Manager"
        if recipient and "name" in recipient:
            salutation_name = recipient["name"]
        
        salutation_para = doc.add_paragraph(f"Dear {salutation_name},")
        salutation_para.paragraph_format.space_after = cfg["space_after_body"]

        # ====================================================================
        # Body Paragraphs
        # ====================================================================
        for paragraph in body_paragraphs:
            body_para = doc.add_paragraph(paragraph)
            body_para.paragraph_format.space_after = cfg["space_after_body"]
            body_para.paragraph_format.line_spacing = cfg["line_spacing"]
            body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in body_para.runs:
                run.font.size = cfg["body_size"]

        # ====================================================================
        # Closing & Signature
        # ====================================================================
        closing_para = doc.add_paragraph("Sincerely,")
        closing_para.runs[0].font.size = cfg["closing_size"]
        closing_para.paragraph_format.space_after = cfg["space_after_closing"]

        # Printed Name
        printed_name_para = doc.add_paragraph(name)
        printed_name_para.runs[0].font.size = cfg["closing_size"]
        printed_name_para.runs[0].font.bold = True
        
        # Simulated Signature (Cursive-ish)
        # We add this BELOW the printed name based on the image style "Olivia Wilson" (signature)
        # Wait, the image has "Sincerely," then "Olivia Wilson" (Printed), then "Olivia Wilson" (Script)
        
        sig_para = doc.add_paragraph()
        sig_run = sig_para.add_run(name)
        sig_run.font.size = cfg["signature_size"]
        # Try to use a script font if available, otherwise italic
        sig_run.font.name = "Brush Script MT" 
        sig_run.font.italic = True

        doc.save(output_path)
        return f"Cover letter successfully created at: {output_path}"

    except Exception as e:
        raise Exception(f"Error creating cover letter: {str(e)}")


@mcp.tool()
def create_formatted_document(
    output_path: str, sections: list, title: str = None
) -> str:
    """
    Create a custom formatted Word document
    """
    try:
        doc = Document()

        if title:
            title_heading = doc.add_heading(title, 0)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for section in sections:
            if "heading" in section:
                doc.add_heading(section["heading"], 1)

            for item in section.get("content", []):
                item_type = item.get("type", "paragraph")
                text = item.get("text", "")

                if item_type == "paragraph":
                    doc.add_paragraph(text)
                elif item_type == "bullet":
                    doc.add_paragraph(text, style="List Bullet")
                elif item_type == "numbered":
                    doc.add_paragraph(text, style="List Number")

            doc.add_paragraph()

        doc.save(output_path)
        return f"Document successfully created at: {output_path}"

    except Exception as e:
        raise Exception(f"Error creating document: {str(e)}")


if __name__ == "__main__":
    try:
        print("Starting docx-creator MCP server...", file=sys.stderr)
        mcp.run()
    except Exception as e:
        print(f"ERROR: Failed to start server: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)

