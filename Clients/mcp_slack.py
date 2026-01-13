"""
Job Search Assistant - Slack Version (Local Storage)
Converted from Streamlit to work with Slack's messaging interface
Uses in-memory storage for local testing
Files are returned in chat, not saved locally
"""

import os
import json
import asyncio
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Optional, List
import uuid
import io
import tempfile

# Slack stuff
from slack_bolt import App
from slack_bolt.adapter.socket_mode import SocketModeHandler
from slack_sdk.errors import SlackApiError

# LLM and agent setup
from langchain_openai import AzureChatOpenAI, ChatOpenAI
from langchain_mcp_adapters.client import MultiServerMCPClient
from langgraph.prebuilt import create_react_agent
from langchain_mcp_adapters.tools import load_mcp_tools
from langgraph.store.memory import InMemoryStore
from langmem import create_manage_memory_tool, create_search_memory_tool

# For reading resumes
import PyPDF2
import docx
from dotenv import load_dotenv
import threading

# ============================================================================
# Basic setup
# ============================================================================

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()

CONFIG_PATH = Path(__file__).resolve().parent / "../Servers/servers_config_slack.json"
CAPABILITIES_PATH = Path(__file__).resolve().parent / "../Servers/server_capabilities.json"

# Only keep the last few messages in context to avoid token limits
WINDOW_TURNS = 3

# ============================================================================
# In-memory user session management - simple storage for local testing
# ============================================================================

class LocalSessionManager:
    """Handles storing user data in memory (no database needed for local testing)"""
    
    def __init__(self):
        # Just store everything in memory
        self.chat_histories = {}  # user_id -> list of messages
        self.user_resumes = {}    # user_id -> resume_text only (no file saving)
    
    def get_chat_history(self, user_id: str) -> List[Dict]:
        """Pull up the conversation history for this user"""
        return self.chat_histories.get(user_id, [])
    
    def add_to_chat_history(self, user_id: str, role: str, content: str):
        """Save a new message to the chat history"""
        if user_id not in self.chat_histories:
            self.chat_histories[user_id] = []
        
        self.chat_histories[user_id].append({
            "role": role,
            "content": content
        })
    
    def get_resume_text(self, user_id: str) -> Optional[str]:
        """Get the user's resume text"""
        return self.user_resumes.get(user_id)
    
    def save_resume_text(self, user_id: str, resume_text: str):
        """Store just the resume text (no file saving)"""
        self.user_resumes[user_id] = resume_text

# ============================================================================
# Resume text extraction - same logic as the Streamlit version
# ============================================================================

def extract_text_from_pdf(pdf_content) -> Optional[str]:
    """Pull text out of a PDF file"""
    try:
        if isinstance(pdf_content, bytes):
            pdf_file = io.BytesIO(pdf_content)
        else:
            pdf_file = open(pdf_content, 'rb')
        
        text = ""
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        return None

def extract_text_from_docx(docx_content) -> Optional[str]:
    """Pull text out of a Word doc"""
    try:
        if isinstance(docx_content, bytes):
            doc = docx.Document(io.BytesIO(docx_content))
        else:
            doc = docx.Document(docx_content)
        
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        return None

def extract_resume_text(file_content, filename: str) -> Optional[str]:
    """Figure out what type of file it is and extract the text"""
    ext = Path(filename).suffix.lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_content)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_content)
    else:
        return None

# ============================================================================
# System prompt - this is the personality and instructions for the agent
# ============================================================================

def load_server_capabilities():
    """Load info about what our MCP servers can do"""
    try:
        with open(CAPABILITIES_PATH) as f:
            return json.load(f)
    except FileNotFoundError:
        logging.warning(f"Server capabilities file not found: {CAPABILITIES_PATH}")
        return {}

def build_enhanced_system_prompt(resume_text: Optional[str] = None) -> str:
    """Build the instructions for the AI agent, including resume context if available"""
    capabilities = load_server_capabilities()
    current_date = datetime.now().strftime("%B %d, %Y")

    base_prompt = f"""You are a Job Search Assistant helping candidates find opportunities and navigate applications. 

IMPORTANT: When creating resumes or cover letters, you MUST use the create_resume or create_cover_letter tools. 
These tools will return the file content that will be automatically shared with the user in Slack.
DO NOT try to save files locally or provide file paths - the MCP server handles file creation and the Slack bot handles delivery.

Today is {current_date}.

## YOUR ROLE & PHILOSOPHY

You help candidates pursue their CAREER GOALS, not just roles matching their current experience. A candidate's current position (e.g., intern, entry-level) does NOT define what they're capable of or aspiring to achieve. Always ask about:
- What roles they WANT to pursue
- What industries or companies interest them
- What skills they want to use or develop
- Their career trajectory and goals

NEVER assume someone wants jobs similar to their current role. An intern may be seeking full-time positions, a data analyst may want to move into engineering, etc.

## WORKFLOW

### 1. DISCOVERY PHASE
Before searching for jobs, understand:
- What TYPE of role are they targeting? (e.g., "Data Scientist", "Software Engineer", "Product Manager")
- What LEVEL? (Intern, Entry-level, Mid-level, Senior)
- Preferred LOCATION or remote preference

Ask clarifying questions! Don't make assumptions.

### 2. JOB SEARCH PHASE
Use the job search tools to find positions matching their GOALS (not just experience):
- Search by their TARGET role title, not current title
- Consider various related titles (e.g., "Data Scientist", "ML Engineer", "Applied Scientist")
- Search across multiple locations if they're flexible
- Cast a wide net initially, then refine based on feedback

Present findings clearly:
- Job title and company
- Location and work arrangement
- Key requirements and responsibilities
- Why it matches their goals
- Any gaps or stretch requirements to address

### 3. APPLICATION STRATEGY PHASE
For jobs they want to apply to:
- Analyze the job description thoroughly
- Identify key requirements and desired qualifications
- Map their experience and skills to requirements
- Suggest how to position their background
- Note any skills to emphasize or gaps to address

### 4. DOCUMENT CREATION PHASE
When creating resumes or cover letters:

**RESUMES:**
- Use the create_resume tool with the resume content
- Tailor to the SPECIFIC job posting
- Lead with relevant skills and projects, not chronological history
- Quantify achievements wherever possible
- Highlight transferable skills from different contexts
- Position current/past roles in terms of relevant skills gained
- Use keywords from the job description naturally
- Keep to ONE page unless explicitly requested otherwise
- The tool will return the file which will be automatically shared in Slack

**COVER LETTERS:**
- Use the create_cover_letter tool with the cover letter content
- Address specific job requirements and company
- Tell the story of WHY they're pursuing this role
- Connect their background to the role's needs (even if indirect)
- Show genuine interest and research about the company
- Address any career transitions or non-traditional paths proactively
- 3-4 substantial paragraphs
- Professional, enthusiastic tone
- The tool will return the file which will be automatically shared in Slack

## KEY PRINCIPLES

1. **Goal-Oriented, Not Experience-Limited**: Help candidates reach for roles they ASPIRE to, not just what they've done
2. **Strategic Positioning**: Frame experience in terms of skills and impact relevant to target role
3. **Proactive Gap Addressing**: Help candidates address experience gaps confidently
4. **Customization is Key**: Every resume and cover letter should be tailored to the specific opportunity
5. **Realistic but Optimistic**: Be honest about stretches while encouraging qualified candidates
6. **Continuous Refinement**: Iterate on documents based on feedback

## COMMUNICATION STYLE

- Ask clarifying questions before taking action
- Explain your reasoning and suggestions
- Offer options when multiple approaches exist
- Be encouraging about career transitions and growth
- Use clear, professional language
- Confirm understanding before creating documents
"""

    if resume_text:
        base_prompt += f"""

CANDIDATE RESUME:
{resume_text}

"""

    base_prompt += """AVAILABLE TOOLS:
"""
    for server_name, config in capabilities.items():
        base_prompt += f"\n{server_name}: {config.get('description', '')}\n"

    return base_prompt

# ============================================================================
# LLM and embedding setup
# ============================================================================

def get_embedding_index():
    """
    Configure the embedding model for memory search
    Tries Azure first, falls back to regular OpenAI if that's not set up
    """
    azure_vars = [
        "AZURE_OPENAI_API_KEY",
        "AZURE_OPENAI_ENDPOINT",
        "AZURE_OPENAI_API_VERSION",
        "AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT",
    ]
    if all(os.getenv(var) for var in azure_vars):
        return {
            "dims": 1536,
            "embed": f"azure_openai:{os.getenv('AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT')}",
        }

    if os.getenv("OPENAI_API_KEY"):
        return {
            "dims": 1536,
            "embed": "openai:text-embedding-3-small",
        }

    raise ValueError(
        f"Missing both Azure and OpenAI embedding configuration. Need either: "
        f"{', '.join(azure_vars)} or OPENAI_API_KEY"
    )

def build_llm():
    """Set up the language model - checks Azure first, then OpenAI"""
    azure_vars = [
        "AZURE_OPENAI_API_KEY",
        "AZURE_OPENAI_ENDPOINT",
        "AZURE_OPENAI_API_VERSION",
        "AZURE_OPENAI_DEPLOYMENT",
    ]
    if all(os.getenv(var) for var in azure_vars):
        return AzureChatOpenAI(
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
            deployment_name=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            temperature=0.1,
            max_tokens=1500,
            timeout=30,
        )

    if os.getenv("OPENAI_API_KEY"):
        return ChatOpenAI(
            api_key=os.getenv("OPENAI_API_KEY"),
            model="gpt-4o",
            temperature=0.1,
            max_tokens=1500,
            timeout=30,
        )

    raise ValueError(
        "Missing both Azure and OpenAI configuration. Need either "
        "AZURE_OPENAI_* vars or OPENAI_API_KEY"
    )

# ============================================================================
# Agent initialization - this is the tricky part with async
# ============================================================================

def get_event_loop():
    """
    Create a persistent event loop that runs in a separate thread
    This is needed because the MCP servers are async but we need sync handling
    """
    loop = asyncio.new_event_loop()

    def run_loop():
        asyncio.set_event_loop(loop)
        loop.run_forever()

    thread = threading.Thread(target=run_loop, daemon=True)
    thread.start()
    return loop

def initialize_memory_system():
    """Set up the memory storage so the agent can remember things"""
    index_cfg = get_embedding_index()
    store = InMemoryStore(index=index_cfg)
    namespace = ("agent_memories",)
    memory_tools = [
        create_manage_memory_tool(namespace, store=store),
        create_search_memory_tool(namespace, store=store),
    ]
    return store, memory_tools

def initialize_agent(loop):
    """
    Initialize the agent with all the MCP server connections
    This takes a while on first run because it's connecting to all the servers
    """
    async def _create_agent():
        try:
            # Load the server config
            with open(CONFIG_PATH) as f:
                config = json.load(f)

            # Handle filesystem root path if it's configured
            filesystem_cfg = config.get("mcpServers", {}).get("filesystem", {})
            fs_args = filesystem_cfg.get("args") if filesystem_cfg else None
            if fs_args:
                env_root = os.getenv("FILESYSTEM_ROOT")
                if env_root:
                    fs_root = Path(env_root).expanduser().resolve()
                else:
                    candidate = Path(fs_args[-1])
                    if not candidate.is_absolute():
                        candidate = (CONFIG_PATH.resolve().parent / candidate).resolve()
                    fs_root = candidate
                fs_args[-1] = str(fs_root)

            servers = config.get("mcpServers", config.get("servers", {}))
            if not servers:
                raise ValueError("No servers found in configuration")

            # Connect to all the MCP servers
            client = MultiServerMCPClient(servers)

            all_tools = []
            sessions = {}

            for server_name in servers.keys():
                try:
                    session_ctx = client.session(server_name)
                    session = await session_ctx.__aenter__()
                    sessions[server_name] = (session, session_ctx)

                    tools = await load_mcp_tools(session)
                    all_tools.extend(tools)
                    logging.info(f"Connected to server: {server_name}, loaded {len(tools)} tools")
                except Exception as e:
                    logging.error(f"Failed to connect to server {server_name}: {e}")
                    continue

            if not all_tools:
                raise ValueError("No tools were loaded from any server")

            # Add memory tools so the agent can remember things
            store, memory_tools = initialize_memory_system()
            all_tools.extend(memory_tools)

            # Set up the LLM
            llm = build_llm()

            # Create the actual agent
            agent = create_react_agent(
                llm,
                all_tools,
                store=store
            )
            logging.info(f"Agent initialized successfully with {len(all_tools)} tools and LangMem memory")
            return agent, client, sessions, store

        except Exception as e:
            logging.error(f"Failed to initialize agent: {e}")
            raise

    future = asyncio.run_coroutine_threadsafe(_create_agent(), loop)
    return future.result(timeout=120)

def run_agent_sync(messages, agent, loop, thread_id: str):
    """Run the agent and wait for a response - wraps the async call to make it synchronous"""
    async def _run_agent():
        try:
            result = await agent.ainvoke(
                {"messages": messages},
                {
                    "recursion_limit": 20,
                    "configurable": {"thread_id": thread_id},
                },
            )
            # Return both the text response and the full result for tool call inspection
            return result
        except Exception as e:
            logging.exception("Agent error")
            return {"error": str(e)}

    future = asyncio.run_coroutine_threadsafe(_run_agent(), loop)
    return future.result(timeout=120)

# ============================================================================
# Start everything up
# ============================================================================

logger.info("Initializing components...")
session_manager = LocalSessionManager()
event_loop = get_event_loop()

logger.info("Initializing agent and connecting to job search servers...")
agent, client, sessions, store = initialize_agent(event_loop)
logger.info("Agent initialized successfully!")

# ============================================================================
# Set up the Slack app
# ============================================================================

app = App(
    token=os.getenv("SLACK_BOT_TOKEN"),
    signing_secret=os.getenv("SLACK_SIGNING_SECRET")
)

# ============================================================================
# Helper function to detect and handle file returns from MCP tools
# ============================================================================

def extract_files_from_tool_calls(result):
    """
    Check if the agent's tool calls returned any files (like .docx documents)
    Returns a list of (filename, file_content) tuples
    """
    files = []
    
    if not result.get("messages"):
        return files
    
    # Look through all messages for tool responses
    for msg in result["messages"]:
        # Check if this is a tool message with content
        if hasattr(msg, 'type') and msg.type == 'tool':
            tool_content = msg.content
            
            # If the tool returned a file path or reference, we need to handle it
            # The docx MCP server should return the file content or a path
            if isinstance(tool_content, str):
                # Check if this looks like a file path
                if tool_content.endswith('.docx'):
                    # Extract filename from path
                    filename = Path(tool_content).name
                    
                    # Try to read the file if it exists
                    if Path(tool_content).exists():
                        with open(tool_content, 'rb') as f:
                            files.append((filename, f.read()))
                        # Delete the file after reading so it doesn't stay on disk
                        try:
                            Path(tool_content).unlink()
                        except Exception as e:
                            logger.warning(f"Could not delete temporary file {tool_content}: {e}")
            
            # Handle base64 or binary content if the tool returns it directly
            elif isinstance(tool_content, bytes):
                # Generate a reasonable filename
                filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                files.append((filename, tool_content))
    
    return files

# ============================================================================
# Slack event handlers
# ============================================================================

@app.event("message")
def handle_message(event, say, client):
    """Handle incoming messages from users"""
    # Ignore messages from bots and threaded replies
    if event.get("subtype") or event.get("bot_id") or event.get("thread_ts"):
        return
    
    user_id = event["user"]
    text = event["text"]
    channel = event["channel"]
    
    # Only respond in DMs or when the bot is mentioned
    channel_type = event.get("channel_type")
    bot_user_id = client.auth_test()['user_id']
    if channel_type != "im" and f"<@{bot_user_id}>" not in text:
        return
    
    # Clean up the bot mention if it's there
    text = text.replace(f"<@{bot_user_id}>", "").strip()
    
    # Process the message
    process_user_input(user_id, text, channel, say, client)

def process_user_input(user_id: str, user_input: str, channel: str, say, slack_client):
    """Handle a user's message and get a response from the agent"""
    # Save the user's message
    session_manager.add_to_chat_history(user_id, "user", user_input)
    
    # Get their conversation history
    msgs = session_manager.get_chat_history(user_id)
    
    # Only use the last few messages to avoid token limits
    window_size = max(2 * WINDOW_TURNS, 2)
    window = msgs[-window_size:]
    
    # Let them know we're working on it
    say("Processing... :hourglass_flowing_sand:")
    
    try:
        # Get their resume if they've uploaded one
        resume_text = session_manager.get_resume_text(user_id)
        
        # Build the full prompt with their resume context
        system_prompt = build_enhanced_system_prompt(resume_text)
        messages = [{"role": "system", "content": system_prompt}] + window
        
        # Run the agent
        result = run_agent_sync(messages, agent, event_loop, thread_id=user_id)
        
        if "error" in result:
            say(f"Error: {result['error']}")
            return
        
        # Get the text response
        reply = result["messages"][-1].content if result.get("messages") else "No reply."
        
        # Check if any files were created
        files = extract_files_from_tool_calls(result)
        
        # Save the response
        session_manager.add_to_chat_history(user_id, "assistant", reply)
        
        # Send the text response
        say(reply)
        
        # Upload any files that were created
        for filename, file_content in files:
            try:
                slack_client.files_upload_v2(
                    channel=channel,
                    file=file_content,
                    filename=filename,
                    title=filename,
                    initial_comment=f"Here's your {filename}"
                )
                logger.info(f"Uploaded file {filename} to channel {channel}")
            except Exception as e:
                logger.error(f"Failed to upload file {filename}: {e}")
                say(f"⚠️ Created {filename} but couldn't upload it to Slack: {str(e)}")
        
    except Exception as e:
        err = f"Error: {str(e)}"
        session_manager.add_to_chat_history(user_id, "assistant", err)
        say(err)

@app.event("file_shared")
def handle_file_upload(event, client, say):
    """Handle when someone uploads a resume"""
    try:
        file_id = event["file_id"]
        user_id = event["user_id"]
        
        # Get info about the file
        file_info = client.files_info(file=file_id)
        file_data = file_info["file"]
        filename = file_data["name"]
        
        # Make sure it's a resume file type we can handle
        if not any(filename.lower().endswith(ext) for ext in ['.pdf', '.docx', '.doc']):
            say("Please upload a PDF or Word document.")
            return
        
        # Download the file from Slack
        file_url = file_data["url_private"]
        
        import requests
        headers = {"Authorization": f"Bearer {os.getenv('SLACK_BOT_TOKEN')}"}
        file_response = requests.get(file_url, headers=headers)
        file_content = file_response.content
        
        # Extract the text from it
        resume_text = extract_resume_text(file_content, filename)
        
        if resume_text:
            # Store ONLY the text, not the file
            session_manager.save_resume_text(user_id, resume_text)
            
            say("✅ Resume uploaded successfully! I'll use this info to help with your job search.")
        else:
            say("Could not extract text from resume. Please make sure it's a valid PDF or Word document.")
            
    except Exception as e:
        logger.exception("Error handling file upload")
        say(f"Error processing file: {str(e)}")

# ============================================================================
# Start the bot
# ============================================================================

if __name__ == "__main__":
    logger.info("Starting Job Search Assistant bot...")
    
    # Check if we have the Socket Mode token
    socket_token = os.getenv("SLACK_APP_TOKEN")
    
    if socket_token:
        # Socket Mode is easier - no need for public URLs or SSL certs
        handler = SocketModeHandler(app, socket_token)
        logger.info("Bot is running! (Socket Mode)")
        handler.start()
    else:
        # HTTP Mode for production - needs a public URL
        from flask import Flask, request
        from slack_bolt.adapter.flask import SlackRequestHandler
        
        flask_app = Flask(__name__)
        handler = SlackRequestHandler(app)
        
        @flask_app.route("/slack/events", methods=["POST"])
        def slack_events():
            return handler.handle(request)
        
        @flask_app.route("/health", methods=["GET"])
        def health():
            return "OK", 200
        
        logger.info("Starting in HTTP Mode on port 3000...")
        flask_app.run(host="0.0.0.0", port=3000)